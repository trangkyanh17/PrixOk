from __future__ import annotations

import asyncio
import json
import os
import re
import time
import uuid
import zipfile
from pathlib import Path
from typing import Any


ARTIFACT_DIR = Path("/app/atri_data/document_artifacts")
MAX_SPEC_BYTES = 1_000_000
MAX_BLOCKS = 240
MAX_TOTAL_TEXT = 300_000
MAX_TABLE_ROWS = 500
MAX_TABLE_COLUMNS = 50
MAX_SHEETS = 12
MAX_SHEET_ROWS = 5_000
MAX_SHEET_COLUMNS = 100
MAX_WORKBOOK_CELLS = 120_000
MAX_ARTIFACTS = 30

# ATRI_DOCUMENT_MODEL_PAYLOAD_COMPAT_V131
_BEGIN_RE = re.compile(r"```\s*atri-document\s*", re.IGNORECASE)
_CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_FILENAME_RE = re.compile(r"[^0-9A-Za-z._-]+")
_SHEET_BAD_RE = re.compile(r"[\\/*?:\[\]]")


class DocumentBridgeError(RuntimeError):
    pass


def _plain_text(value: Any, limit: int = 50_000) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list, tuple)):
        value = json.dumps(value, ensure_ascii=False)
    text = _CONTROL_RE.sub("", str(value)).strip()
    return text[:limit]


def strip_document_envelope(text: str) -> str:
    cleaned = str(text or "")
    while True:
        match = _BEGIN_RE.search(cleaned)
        if match is None:
            return cleaned.strip()
        end = cleaned.find("```", match.end())
        if end < 0:
            return cleaned[: match.start()].rstrip()
        cleaned = cleaned[: match.start()] + cleaned[end + 3 :]


def extract_document_spec(
    text: str,
    *,
    strict: bool = True,
) -> tuple[str, dict[str, Any] | None, str | None]:
    raw = str(text or "")
    match = _BEGIN_RE.search(raw)
    if match is None:
        return raw.strip(), None, None

    end = raw.find("```", match.end())
    if end < 0:
        cleaned = raw[: match.start()].rstrip()
        if strict:
            raise DocumentBridgeError("DOCUMENT_SPEC_FENCE_UNTERMINATED")
        return cleaned, None, "DOCUMENT_SPEC_FENCE_UNTERMINATED"

    payload = raw[match.end() : end].strip()
    cleaned = (raw[: match.start()] + raw[end + 3 :]).strip()

    if _BEGIN_RE.search(cleaned) is not None:
        cleaned = strip_document_envelope(cleaned)
        if strict:
            raise DocumentBridgeError("DOCUMENT_SPEC_MULTIPLE_ENVELOPES")
        return cleaned, None, "DOCUMENT_SPEC_MULTIPLE_ENVELOPES"

    if len(payload.encode("utf-8")) > MAX_SPEC_BYTES:
        if strict:
            raise DocumentBridgeError("DOCUMENT_SPEC_TOO_LARGE")
        return cleaned, None, "DOCUMENT_SPEC_TOO_LARGE"

    try:
        decoded = json.loads(payload)
    except Exception as exc:
        if strict:
            raise DocumentBridgeError(
                "DOCUMENT_SPEC_JSON_INVALID:" + type(exc).__name__
            ) from exc
        return cleaned, None, "DOCUMENT_SPEC_JSON_INVALID"

    if not isinstance(decoded, dict):
        if strict:
            raise DocumentBridgeError("DOCUMENT_SPEC_NOT_OBJECT")
        return cleaned, None, "DOCUMENT_SPEC_NOT_OBJECT"

    return cleaned, decoded, None


def _normalize_format(spec: dict[str, Any]) -> str:
    value = _plain_text(spec.get("format"), 20).lower().lstrip(".")
    if not value:
        filename = Path(
            _plain_text(spec.get("filename"), 160)
        ).name
        value = Path(filename).suffix.lower().lstrip(".")
    aliases = {
        "word": "docx",
        "excel": "xlsx",
        "spreadsheet": "xlsx",
    }
    value = aliases.get(value, value)
    if value not in {"pdf", "docx", "xlsx"}:
        raise DocumentBridgeError("DOCUMENT_FORMAT_NOT_ALLOWED")
    return value


def _safe_filename(value: Any, fmt: str) -> str:
    raw = Path(_plain_text(value, 160) or f"atri-document.{fmt}").name
    stem = Path(raw).stem[:72]
    stem = _FILENAME_RE.sub("-", stem).strip(".-_") or "atri-document"
    return f"{stem}.{fmt}"


def _artifact_path(filename: str) -> Path:
    ARTIFACT_DIR.mkdir(parents=True, exist_ok=True, mode=0o700)
    os.chmod(ARTIFACT_DIR, 0o700)
    stamp = time.strftime("%Y%m%d-%H%M%S")
    token = uuid.uuid4().hex[:8]
    path = ARTIFACT_DIR / f"{stamp}-{token}-{filename}"
    resolved = path.resolve()
    if ARTIFACT_DIR.resolve() not in resolved.parents:
        raise DocumentBridgeError("DOCUMENT_ARTIFACT_PATH_ESCAPE")
    return resolved


def _cleanup_artifacts() -> None:
    try:
        ARTIFACT_DIR.mkdir(parents=True, exist_ok=True, mode=0o700)
        files = sorted(
            (item for item in ARTIFACT_DIR.iterdir() if item.is_file()),
            key=lambda item: item.stat().st_mtime,
            reverse=True,
        )
        for item in files[MAX_ARTIFACTS:]:
            item.unlink(missing_ok=True)
    except Exception:
        pass


def _normalize_blocks(spec: dict[str, Any]) -> list[dict[str, Any]]:
    source = spec.get("blocks")
    content_source = spec.get("content")
    if (
        not isinstance(source, list)
        or not source
    ) and isinstance(content_source, list):
        source = content_source
    blocks: list[dict[str, Any]] = []

    if isinstance(source, list):
        for item in source[:MAX_BLOCKS]:
            if isinstance(item, dict):
                blocks.append(dict(item))
            else:
                blocks.append({"type": "paragraph", "text": item})

    if not blocks:
        content = _plain_text(spec.get("content"), MAX_TOTAL_TEXT)
        for paragraph in re.split(r"\n\s*\n", content):
            paragraph = paragraph.strip()
            if paragraph:
                blocks.append({"type": "paragraph", "text": paragraph})

    total = 0
    normalized: list[dict[str, Any]] = []
    for block in blocks:
        kind = _plain_text(block.get("type"), 30).lower() or "paragraph"
        if kind not in {
            "heading",
            "paragraph",
            "bullet",
            "numbered",
            "table",
            "page_break",
        }:
            kind = "paragraph"

        current: dict[str, Any] = {"type": kind}
        if kind == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            if not isinstance(headers, list):
                headers = []
            if not isinstance(rows, list):
                rows = []
            safe_headers = [
                _plain_text(value, 2_000)
                for value in headers[:MAX_TABLE_COLUMNS]
            ]
            safe_rows: list[list[str]] = []
            for row in rows[:MAX_TABLE_ROWS]:
                if not isinstance(row, list):
                    row = [row]
                safe_rows.append(
                    [
                        _plain_text(value, 2_000)
                        for value in row[:MAX_TABLE_COLUMNS]
                    ]
                )
            current["headers"] = safe_headers
            current["rows"] = safe_rows
            total += sum(len(value) for value in safe_headers)
            total += sum(len(value) for row in safe_rows for value in row)
        elif kind == "heading":
            current["text"] = _plain_text(block.get("text"), 20_000)
            try:
                level = int(block.get("level", 1))
            except Exception:
                level = 1
            current["level"] = max(1, min(level, 4))
            total += len(current["text"])
        elif kind != "page_break":
            current["text"] = _plain_text(block.get("text"), 50_000)
            total += len(current["text"])

        if total > MAX_TOTAL_TEXT:
            raise DocumentBridgeError("DOCUMENT_TEXT_LIMIT_EXCEEDED")
        normalized.append(current)

    if not normalized:
        raise DocumentBridgeError("DOCUMENT_CONTENT_EMPTY")
    return normalized


def _create_docx(
    path: Path,
    title: str,
    blocks: list[dict[str, Any]],
) -> dict[str, Any]:
    from docx import Document
    from docx.shared import Inches

    document = Document()
    for section in document.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    if title:
        document.add_heading(title, level=0)

    for block in blocks:
        kind = block["type"]
        if kind == "heading":
            document.add_heading(block.get("text", ""), level=block.get("level", 1))
        elif kind == "paragraph":
            document.add_paragraph(block.get("text", ""))
        elif kind == "bullet":
            document.add_paragraph(block.get("text", ""), style="List Bullet")
        elif kind == "numbered":
            document.add_paragraph(block.get("text", ""), style="List Number")
        elif kind == "page_break":
            document.add_page_break()
        elif kind == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            width = max(
                [len(headers)] + [len(row) for row in rows] + [1]
            )
            table = document.add_table(rows=0, cols=width)
            table.style = "Table Grid"
            if headers:
                cells = table.add_row().cells
                for index, value in enumerate(headers[:width]):
                    cells[index].text = value
            for row in rows:
                cells = table.add_row().cells
                for index, value in enumerate(row[:width]):
                    cells[index].text = value

    document.save(path)

    with zipfile.ZipFile(path) as package:
        if package.testzip() is not None:
            raise DocumentBridgeError("DOCX_ZIP_VALIDATION_FAILED")
    reopened = Document(path)
    if not reopened.paragraphs and not reopened.tables:
        raise DocumentBridgeError("DOCX_REOPEN_EMPTY")
    return {
        "format": "docx",
        "paragraphs": len(reopened.paragraphs),
        "tables": len(reopened.tables),
    }


def _find_pdf_font() -> Path | None:
    candidates = (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
        "/usr/share/fonts/opentype/noto/NotoSans-Regular.ttf",
        "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
    )
    for value in candidates:
        path = Path(value)
        if path.is_file():
            return path
    return None


def _wrap_text(text: str, width: int) -> list[str]:
    lines: list[str] = []
    for paragraph in str(text or "").splitlines() or [""]:
        words = paragraph.split()
        if not words:
            lines.append("")
            continue
        current = ""
        for word in words:
            while len(word) > width:
                head, word = word[:width], word[width:]
                if current:
                    lines.append(current)
                    current = ""
                lines.append(head)
            candidate = word if not current else current + " " + word
            if len(candidate) <= width:
                current = candidate
            else:
                lines.append(current)
                current = word
        if current:
            lines.append(current)
    return lines


def _create_pdf(
    path: Path,
    title: str,
    blocks: list[dict[str, Any]],
) -> dict[str, Any]:
    try:
        import pymupdf as fitz
    except ImportError:
        import fitz  # type: ignore[no-redef]

    document = fitz.open()
    font_path = _find_pdf_font()
    page = None
    y = 0.0
    page_width = 595.0
    page_height = 842.0
    margin = 48.0

    def new_page():
        nonlocal page, y
        page = document.new_page(width=page_width, height=page_height)
        if font_path is not None:
            page.insert_font(fontname="AtriFont", fontfile=str(font_path))
        y = margin

    def write_lines(text: str, size: float, *, indent: float = 0.0, after: float = 5.0):
        nonlocal y
        if page is None:
            new_page()
        font_name = "AtriFont" if font_path is not None else "helv"
        width = max(25, int((page_width - 2 * margin - indent) / max(size * 0.52, 1)))
        for line in _wrap_text(text, width):
            line_height = size * 1.45
            if y + line_height > page_height - margin:
                new_page()
            if line:
                page.insert_text(
                    (margin + indent, y + size),
                    line,
                    fontname=font_name,
                    fontsize=size,
                    color=(0, 0, 0),
                )
            y += line_height
        y += after

    if title:
        write_lines(title, 18.0, after=10.0)

    for block in blocks:
        kind = block["type"]
        if kind == "page_break":
            new_page()
        elif kind == "heading":
            level = block.get("level", 1)
            size = {1: 16.0, 2: 14.0, 3: 12.5, 4: 11.5}.get(level, 12.0)
            write_lines(block.get("text", ""), size, after=7.0)
        elif kind == "paragraph":
            write_lines(block.get("text", ""), 10.5)
        elif kind == "bullet":
            write_lines("- " + block.get("text", ""), 10.5, indent=12.0)
        elif kind == "numbered":
            write_lines("1. " + block.get("text", ""), 10.5, indent=12.0)
        elif kind == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            if headers:
                write_lines(" | ".join(headers), 9.5, after=2.0)
            for row in rows:
                write_lines(" | ".join(row), 9.0, after=1.0)
            y += 5.0

    if document.page_count == 0:
        new_page()
    document.set_metadata({"title": title or path.stem, "producer": "Atri V1.2.8"})
    document.save(path, garbage=4, deflate=True)
    document.close()

    with fitz.open(path) as reopened:
        page_count = reopened.page_count
        if page_count < 1:
            raise DocumentBridgeError("PDF_REOPEN_EMPTY")
        text_chars = sum(len(page.get_text("text")) for page in reopened)
    return {
        "format": "pdf",
        "pages": page_count,
        "text_chars": text_chars,
    }


def _safe_sheet_name(value: Any, used: set[str], index: int) -> str:
    name = _SHEET_BAD_RE.sub("_", _plain_text(value, 80)).strip("'")[:31]
    name = name or f"Sheet{index}"
    base = name
    counter = 2
    while name.casefold() in used:
        suffix = f"-{counter}"
        name = base[: 31 - len(suffix)] + suffix
        counter += 1
    used.add(name.casefold())
    return name


def _cell_value(value: Any) -> Any:
    if value is None or isinstance(value, (str, int, float, bool)):
        return value
    return _plain_text(value, 20_000)


def _normalize_sheets(spec: dict[str, Any]) -> list[dict[str, Any]]:
    source = spec.get("sheets")
    if not isinstance(source, list) or not source:
        blocks = _normalize_blocks(spec)
        rows = [["Type", "Content"]]
        for block in blocks:
            if block["type"] == "table":
                if block.get("headers"):
                    rows.append(block["headers"])
                rows.extend(block.get("rows", []))
            elif block["type"] != "page_break":
                rows.append([block["type"], block.get("text", "")])
        source = [{"name": "Data", "rows": rows}]

    used: set[str] = set()
    sheets: list[dict[str, Any]] = []
    cells = 0
    for index, item in enumerate(source[:MAX_SHEETS], start=1):
        if not isinstance(item, dict):
            raise DocumentBridgeError("XLSX_SHEET_NOT_OBJECT")
        rows = item.get("rows", [])
        if not isinstance(rows, list):
            raise DocumentBridgeError("XLSX_ROWS_NOT_LIST")
        safe_rows: list[list[Any]] = []
        for row in rows[:MAX_SHEET_ROWS]:
            if not isinstance(row, list):
                row = [row]
            safe_row = [_cell_value(value) for value in row[:MAX_SHEET_COLUMNS]]
            cells += len(safe_row)
            if cells > MAX_WORKBOOK_CELLS:
                raise DocumentBridgeError("XLSX_CELL_LIMIT_EXCEEDED")
            safe_rows.append(safe_row)
        sheets.append(
            {
                "name": _safe_sheet_name(item.get("name"), used, index),
                "rows": safe_rows,
                "freeze_panes": _plain_text(item.get("freeze_panes"), 20),
                "auto_filter": bool(item.get("auto_filter", True)),
            }
        )
    if not sheets:
        raise DocumentBridgeError("XLSX_SHEETS_EMPTY")
    return sheets


def _create_xlsx(path: Path, spec: dict[str, Any]) -> dict[str, Any]:
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    sheets = _normalize_sheets(spec)
    workbook = Workbook()
    workbook.remove(workbook.active)

    total_rows = 0
    for sheet_spec in sheets:
        sheet = workbook.create_sheet(title=sheet_spec["name"])
        rows = sheet_spec["rows"]
        for row in rows:
            sheet.append(row)
        total_rows += len(rows)

        if rows:
            for cell in sheet[1]:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor="1F4E78")
                cell.alignment = Alignment(horizontal="center", vertical="center")

        freeze = sheet_spec["freeze_panes"]
        if freeze and re.fullmatch(r"[A-Z]{1,3}[1-9][0-9]{0,5}", freeze):
            sheet.freeze_panes = freeze
        elif rows:
            sheet.freeze_panes = "A2"

        if sheet_spec["auto_filter"] and sheet.max_row >= 1 and sheet.max_column >= 1:
            sheet.auto_filter.ref = sheet.dimensions

        for column in range(1, sheet.max_column + 1):
            longest = 0
            for row in range(1, min(sheet.max_row, 500) + 1):
                value = sheet.cell(row=row, column=column).value
                longest = max(longest, len(str(value or "")))
            sheet.column_dimensions[get_column_letter(column)].width = min(
                max(longest + 2, 10),
                50,
            )

    workbook.save(path)
    reopened = load_workbook(path, read_only=False, data_only=False)
    names = list(reopened.sheetnames)
    if names != [sheet["name"] for sheet in sheets]:
        reopened.close()
        raise DocumentBridgeError("XLSX_REOPEN_SHEET_MISMATCH")
    reopened.close()
    return {
        "format": "xlsx",
        "sheets": len(sheets),
        "rows": total_rows,
    }


def execute_document_spec(spec: dict[str, Any]) -> dict[str, Any]:
    if int(spec.get("version", 1)) != 1:
        raise DocumentBridgeError("DOCUMENT_SPEC_VERSION_NOT_1")

    fmt = _normalize_format(spec)
    filename = _safe_filename(spec.get("filename"), fmt)
    title = _plain_text(spec.get("title"), 2_000)
    path = _artifact_path(filename)

    if fmt == "xlsx":
        details = _create_xlsx(path, spec)
    else:
        blocks = _normalize_blocks(spec)
        if fmt == "docx":
            details = _create_docx(path, title, blocks)
        else:
            details = _create_pdf(path, title, blocks)

    if not path.is_file() or path.stat().st_size <= 100:
        path.unlink(missing_ok=True)
        raise DocumentBridgeError("DOCUMENT_ARTIFACT_INVALID")

    os.chmod(path, 0o600)
    _cleanup_artifacts()
    return {
        "executed": True,
        "format": fmt,
        "filename": filename,
        "artifact_path": str(path),
        "artifact_bytes": path.stat().st_size,
        "details": details,
    }


async def process_document_response(message: Any, response_text: str) -> dict[str, Any]:
    raw = str(response_text or "")
    try:
        clean_text, spec, parse_error = extract_document_spec(raw, strict=False)
        if spec is None:
            if parse_error:
                clean_text = (
                    clean_text
                    + "\n\nKhông thể tạo tệp: đặc tả tài liệu không hợp lệ."
                ).strip()
                return {
                    "executed": False,
                    "clean_text": clean_text,
                    "error": parse_error,
                }
            return {
                "executed": False,
                "clean_text": raw,
                "reason": "no_document_spec",
            }

        result = await asyncio.to_thread(execute_document_spec, spec)
        sender = getattr(message, "reply_document", None)
        if not callable(sender):
            raise DocumentBridgeError("TELEGRAM_REPLY_DOCUMENT_UNAVAILABLE")

        caption = (
            f"{result['format'].upper()} đã được tạo và kiểm tra"
            f" ({result['artifact_bytes']} bytes)."
        )
        await sender(
            document=result["artifact_path"],
            caption=caption,
            quote=True,
        )

        result["telegram_sent"] = True
        result["clean_text"] = clean_text or "Tệp đã được tạo và gửi kèm."
        return result
    except Exception as exc:
        clean_text = strip_document_envelope(raw)
        clean_text = (
            clean_text
            + "\n\nKhông thể tạo tệp: "
            + type(exc).__name__
            + "."
        ).strip()
        return {
            "executed": False,
            "clean_text": clean_text,
            "error": type(exc).__name__ + ":" + str(exc)[:500],
        }
